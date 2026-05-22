from __future__ import annotations

from datetime import datetime, timezone
from io import BytesIO
import os
import re
from pathlib import Path
import tempfile
from typing import Any

import httpx
from docx import Document
from openpyxl import Workbook

# Use a more reliable persistent location in serverless environments
# Try /tmp (common in serverless), fall back to temp directory
_FONT_BASE_DIR = Path(os.environ.get("FONT_CACHE_DIR", "/tmp" if Path("/tmp").exists() else tempfile.gettempdir()))
_KHMER_FONT_PATH = _FONT_BASE_DIR / "NotoSansKhmer-Regular.ttf"
_LATIN_FONT_PATH = _FONT_BASE_DIR / "NotoSans-Regular.ttf"
_KHMER_FONT_URLS = (
    "https://cdn.jsdelivr.net/gh/googlefonts/noto-fonts@latest/hinted/ttf/NotoSansKhmer/NotoSansKhmer-Regular.ttf",
    "https://github.com/notofonts/khmer/raw/main/fonts/ttf/NotoSansKhmer/NotoSansKhmer-Regular.ttf",
    "https://fonts.google.com/download?family=Noto%20Sans%20Khmer",
)
_LATIN_FONT_URLS = (
    "https://cdn.jsdelivr.net/gh/googlefonts/noto-fonts@latest/hinted/ttf/NotoSans/NotoSans-Regular.ttf",
    "https://github.com/notofonts/noto-fonts/raw/main/hinted/ttf/NotoSans/NotoSans-Regular.ttf",
)
_VALID_TTF_HEADERS = {b"\x00\x01\x00\x00", b"OTTO", b"true", b"ttcf"}
_KHMER_RE = re.compile(r"[\u1780-\u17ff]")
_KHMER_CLUSTER_RE = re.compile(
    r"[\u1780-\u17a2][\u17b6-\u17c8\u17cb-\u17d1\u17d3-\u17dd]*(?:\u17d2[\u1780-\u17a2][\u17b6-\u17c8\u17cb-\u17d1\u17d3-\u17dd]*)*"
    r"|[\u17b6-\u17c8\u17cb-\u17d1\u17d3-\u17dd]"
    r"|.",
    re.DOTALL,
)

# Cache font data in memory as fallback
_KHMER_FONT_CACHE: bytes | None = None
_LATIN_FONT_CACHE: bytes | None = None


def _is_valid_font_file(path: Path) -> bool:
    try:
        return path.exists() and path.stat().st_size > 100_000 and path.read_bytes()[:4] in _VALID_TTF_HEADERS
    except Exception:
        return False


async def ensure_khmer_font() -> str | None:
    """Ensure Khmer font is available, trying file system then memory cache."""
    global _KHMER_FONT_CACHE
    
    # Try file system first
    if _is_valid_font_file(_KHMER_FONT_PATH):
        return str(_KHMER_FONT_PATH)

    # Try to download and cache in file system
    for idx, font_url in enumerate(_KHMER_FONT_URLS):
        try:
            async with httpx.AsyncClient(timeout=30, follow_redirects=True) as client:
                response = await client.get(font_url)
                response.raise_for_status()
                font_data = response.content
                
                # Validate font data
                if font_data[:4] not in _VALID_TTF_HEADERS or len(font_data) < 100_000:
                    continue  # Invalid font, try next URL
                
                # Cache in memory for this session
                _KHMER_FONT_CACHE = font_data
                
                # Try to persist to file system if possible
                try:
                    _KHMER_FONT_PATH.parent.mkdir(parents=True, exist_ok=True)
                    _KHMER_FONT_PATH.write_bytes(font_data)
                except Exception:
                    pass  # Continue even if file write fails
                
                return str(_KHMER_FONT_PATH)
        except Exception:
            continue

    # If all downloads failed, check if we have cached font from memory
    if _KHMER_FONT_CACHE and _KHMER_FONT_CACHE[:4] in _VALID_TTF_HEADERS:
        # Write to temp location for this request
        try:
            temp_font = Path(tempfile.gettempdir()) / f"khmer_font_{id(_KHMER_FONT_CACHE)}.ttf"
            temp_font.write_bytes(_KHMER_FONT_CACHE)
            return str(temp_font)
        except Exception:
            pass

    return None


async def ensure_latin_font() -> str | None:
    global _LATIN_FONT_CACHE

    if _is_valid_font_file(_LATIN_FONT_PATH):
        return str(_LATIN_FONT_PATH)

    for font_url in _LATIN_FONT_URLS:
        try:
            async with httpx.AsyncClient(timeout=30, follow_redirects=True) as client:
                response = await client.get(font_url)
                response.raise_for_status()
                font_data = response.content

                if font_data[:4] not in _VALID_TTF_HEADERS or len(font_data) < 100_000:
                    continue

                _LATIN_FONT_CACHE = font_data

                try:
                    _LATIN_FONT_PATH.parent.mkdir(parents=True, exist_ok=True)
                    _LATIN_FONT_PATH.write_bytes(font_data)
                except Exception:
                    pass

                return str(_LATIN_FONT_PATH)
        except Exception:
            continue

    if _LATIN_FONT_CACHE and _LATIN_FONT_CACHE[:4] in _VALID_TTF_HEADERS:
        try:
            temp_font = Path(tempfile.gettempdir()) / f"latin_font_{id(_LATIN_FONT_CACHE)}.ttf"
            temp_font.write_bytes(_LATIN_FONT_CACHE)
            return str(temp_font)
        except Exception:
            pass

    return None


def _build_unicode_pdf() -> Any:
    """Build a new PDF with Khmer font support using proper Unicode handling."""
    from fpdf import FPDF  # noqa: PLC0415

    pdf = FPDF(font_cache_dir="DEPRECATED")
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_margins(left=15, top=15, right=15)
    
    # Add Khmer font with Unicode support
    font_path = str(_KHMER_FONT_PATH) if _KHMER_FONT_PATH.exists() else None
    latin_font_path = str(_LATIN_FONT_PATH) if _LATIN_FONT_PATH.exists() else None
    
    if not font_path and _KHMER_FONT_CACHE:
        try:
            temp_font = Path(tempfile.gettempdir()) / "khmer_font_temp.ttf"
            temp_font.write_bytes(_KHMER_FONT_CACHE)
            font_path = str(temp_font)
        except Exception:
            pass
    
    if not font_path:
        raise RuntimeError("Khmer font not available - could not download or cache font file")
    
    try:
        # Register Khmer font with full Unicode support and text shaping
        pdf.add_font(family="Khmer", fname=font_path, uni=True)
        if latin_font_path:
            pdf.add_font(family="Latin", fname=latin_font_path, uni=True)
            pdf.set_fallback_fonts(["Latin"], exact_match=False)
    except Exception as e:
        raise RuntimeError(f"Failed to add Khmer font: {e}") from e
    
    try:
        pdf.set_text_shaping(True, direction="ltr", script="khmr", language="khm")
    except Exception as e:
        raise RuntimeError("Khmer PDF shaping is unavailable. Install fpdf2[shaping] with uharfbuzz.") from e
    pdf.add_page()
    pdf.set_font("Khmer", size=12)
    
    return pdf


def _has_khmer(text: str) -> bool:
    return bool(_KHMER_RE.search(text or ""))


def _clean_pdf_text(text: str, fallback: str = "") -> str:
    try:
        value = str(text).replace("\r\n", "\n").replace("\r", "\n").strip()
        value.encode("utf-8").decode("utf-8")
        return value or fallback
    except Exception:
        return fallback


def _khmer_clusters(text: str) -> list[str]:
    return [match.group(0) for match in _KHMER_CLUSTER_RE.finditer(text)]


def _break_long_token(pdf: Any, token: str, max_width: float) -> list[str]:
    if pdf.get_string_width(token) <= max_width:
        return [token]

    pieces = _khmer_clusters(token) if _has_khmer(token) else list(token)
    lines: list[str] = []
    current = ""
    for piece in pieces:
        candidate = current + piece
        if not current or pdf.get_string_width(candidate) <= max_width:
            current = candidate
            continue
        lines.append(current)
        current = piece
    if current:
        lines.append(current)
    return lines or [token]


def _wrap_pdf_text(pdf: Any, text: str, max_width: float) -> list[str]:
    lines: list[str] = []
    for paragraph in text.split("\n"):
        if not paragraph.strip():
            lines.append("")
            continue

        tokens = re.findall(r"\S+|\s+", paragraph)
        current = ""
        for token in tokens:
            candidate = current + token
            if not current or pdf.get_string_width(candidate) <= max_width:
                current = candidate
                continue

            if current.strip():
                lines.extend(_break_long_token(pdf, current.rstrip(), max_width))
            current = token.lstrip()

        if current.strip():
            lines.extend(_break_long_token(pdf, current.rstrip(), max_width))

    return lines or [""]


def _write_wrapped_text(pdf: Any, text: str, line_height: float) -> None:
    max_width = pdf.w - pdf.l_margin - pdf.r_margin
    for line in _wrap_pdf_text(pdf, text, max_width):
        pdf.cell(w=0, h=line_height, text=line, new_x="LMARGIN", new_y="NEXT")


def _render_pdf_title(pdf: Any, title: str) -> None:
    title_text = _clean_pdf_text(title, "Chat Export")

    pdf.set_font("Khmer", size=16)
    pdf.set_text_color(0, 0, 0)
    _write_wrapped_text(pdf, title_text, 10)
    pdf.ln(3)


def _render_pdf_generated_at(pdf: Any, khmer: bool) -> None:
    pdf.set_font("Khmer", size=9)
    pdf.set_text_color(120, 120, 120)
    now_str = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    label = "បានបង្កើត" if khmer else "Generated"
    pdf.cell(w=0, h=6, text=f"{label}: {now_str}", new_x="LMARGIN", new_y="NEXT")
    pdf.set_text_color(0, 0, 0)
    pdf.ln(3)


def _render_pdf_message(pdf: Any, speaker: str, content: str, created_at: str = "") -> None:
    timestamp = ""
    if created_at:
        try:
            dt = datetime.fromisoformat(str(created_at).replace("Z", "+00:00"))
            timestamp = dt.strftime("%H:%M")
        except Exception:
            pass

    header = f"{speaker} | {timestamp}" if timestamp else f"{speaker}"

    pdf.set_font("Khmer", size=9)
    pdf.set_text_color(80, 80, 80)
    pdf.cell(w=0, h=6, text=header, new_x="LMARGIN", new_y="NEXT")

    if content and content.strip():
        pdf.set_font("Khmer", size=11)
        pdf.set_text_color(0, 0, 0)
        content_clean = _clean_pdf_text(content, "[Text encoding error]")
        _write_wrapped_text(pdf, content_clean, 7)
    
    pdf.ln(2)


async def build_khmer_pdf_full(
    title: str,
    messages: list[dict[str, Any]],
) -> bytes:
    """Build a PDF with Khmer font support, ensuring font is loaded first."""
    # Ensure font is loaded before building PDF
    font_result = await ensure_khmer_font()
    if not font_result:
        raise RuntimeError("Failed to load Khmer font for PDF generation")
    await ensure_latin_font()
    
    pdf = _build_unicode_pdf()
    is_khmer_export = _has_khmer(title) or any(_has_khmer(str(msg.get("content", ""))) for msg in messages)
    _render_pdf_title(pdf, title)
    _render_pdf_generated_at(pdf, is_khmer_export)

    for msg in messages:
        sender_type = msg.get("sender_type", "user")
        if is_khmer_export:
            speaker = "ជំនួយការ" if sender_type == "assistant" else "អ្នក"
        else:
            speaker = "Bot" if sender_type == "assistant" else "User"
        _render_pdf_message(
            pdf,
            speaker=speaker,
            content=str(msg.get("content", "")),
            created_at=str(msg.get("created_at", "")),
        )

    return bytes(pdf.output())


def _safe_stem(text: str) -> str:
    normalized = re.sub(r"[^a-zA-Z0-9]+", "-", text.strip().lower()).strip("-")
    return normalized[:40] or "saved-exchange"


def build_exchange_text(user_text: str, assistant_text: str) -> str:
    return f"User: {user_text.strip()}\n\nAI: {assistant_text.strip()}\n"


def build_exchange_filename(user_text: str, suffix: str) -> str:
    return f"{_safe_stem(user_text)}.{suffix}"


def build_content_filename(title: str, suffix: str) -> str:
    return f"{_safe_stem(title)}.{suffix}"


def build_text_export(user_text: str, assistant_text: str) -> tuple[str, bytes, str]:
    content = build_exchange_text(user_text, assistant_text).encode("utf-8")
    return (
        build_exchange_filename(user_text, "txt"),
        content,
        "text/plain",
    )


def build_text_content_export(title: str, content: str) -> tuple[str, bytes, str]:
    return (
        build_content_filename(title, "txt"),
        content.strip().encode("utf-8"),
        "text/plain",
    )


def build_docx_export(user_text: str, assistant_text: str) -> tuple[str, bytes, str]:
    document = Document()
    document.add_heading("Saved Exchange", level=1)
    document.add_paragraph("User")
    document.add_paragraph(user_text.strip())
    document.add_paragraph("AI")
    document.add_paragraph(assistant_text.strip())
    buffer = BytesIO()
    document.save(buffer)
    return (
        build_exchange_filename(user_text, "docx"),
        buffer.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def build_docx_content_export(title: str, content: str) -> tuple[str, bytes, str]:
    document = Document()
    document.add_heading(title.strip() or "Saved Conversation", level=1)
    for block in content.strip().split("\n\n"):
        document.add_paragraph(block.strip())
    buffer = BytesIO()
    document.save(buffer)
    return (
        build_content_filename(title, "docx"),
        buffer.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def build_excel_export(user_text: str, assistant_text: str) -> tuple[str, bytes, str]:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Exchange"
    sheet.append(["Speaker", "Message"])
    sheet.append(["User", user_text.strip()])
    sheet.append(["AI", assistant_text.strip()])
    sheet.column_dimensions["A"].width = 16
    sheet.column_dimensions["B"].width = 100
    buffer = BytesIO()
    workbook.save(buffer)
    return (
        build_exchange_filename(user_text, "xlsx"),
        buffer.getvalue(),
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


def build_excel_content_export(title: str, content: str) -> tuple[str, bytes, str]:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Conversation"
    sheet.append(["Title", title.strip() or "Saved Conversation"])
    sheet.append(["Speaker", "Message"])
    for block in content.strip().split("\n\n"):
        if ":" in block:
            speaker, message = block.split(":", 1)
            sheet.append([speaker.strip(), message.strip()])
        else:
            sheet.append(["Note", block.strip()])
    sheet.column_dimensions["A"].width = 16
    sheet.column_dimensions["B"].width = 100
    buffer = BytesIO()
    workbook.save(buffer)
    return (
        build_content_filename(title, "xlsx"),
        buffer.getvalue(),
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


def build_pdf_export(user_text: str, assistant_text: str) -> tuple[str, bytes, str]:
    pdf = _build_unicode_pdf()
    _render_pdf_title(pdf, "Saved Exchange")
    _render_pdf_generated_at(pdf)
    _render_pdf_message(pdf, speaker="User", content=user_text)
    _render_pdf_message(pdf, speaker="AI", content=assistant_text)
    return (
        build_exchange_filename(user_text, "pdf"),
        bytes(pdf.output()),
        "application/pdf",
    )


def build_pdf_content_export(title: str, content: str) -> tuple[str, bytes, str]:
    pdf = _build_unicode_pdf()
    _render_pdf_title(pdf, title.strip() or "Saved Conversation")
    _render_pdf_generated_at(pdf)
    for block in content.strip().split("\n\n"):
        text = block.strip()
        if not text:
            continue
        if ":" in text:
            speaker, message = text.split(":", 1)
            _render_pdf_message(pdf, speaker=speaker.strip() or "Note", content=message.strip())
        else:
            _render_pdf_message(pdf, speaker="Note", content=text)
    return (
        build_content_filename(title, "pdf"),
        bytes(pdf.output()),
        "application/pdf",
    )


def build_exchange_export(file_format: str, user_text: str, assistant_text: str) -> tuple[str, bytes, str]:
    if file_format == "text":
        return build_text_export(user_text, assistant_text)
    if file_format == "pdf":
        return build_pdf_export(user_text, assistant_text)
    if file_format == "word":
        return build_docx_export(user_text, assistant_text)
    if file_format == "excel":
        return build_excel_export(user_text, assistant_text)
    raise ValueError("Unsupported export format.")


def build_content_export(file_format: str, title: str, content: str) -> tuple[str, bytes, str]:
    if file_format == "text":
        return build_text_content_export(title, content)
    if file_format == "pdf":
        return build_pdf_content_export(title, content)
    if file_format == "word":
        return build_docx_content_export(title, content)
    if file_format == "excel":
        return build_excel_content_export(title, content)
    raise ValueError("Unsupported export format.")
